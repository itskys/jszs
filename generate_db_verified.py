import os
import hashlib
import json
import re
try:
    from docx import Document
except ImportError:
    print("❌ 请安装依赖: pip install python-docx")
    exit()

# ================= 配置 =================
INPUT_DIR = './2025tk'
OUTPUT_JS = 'questions_data.js'
REPORT_FILE = 'extraction_report.txt'

# ================= 正则 =================
RE_OPTION = re.compile(r'^\s*\(?([A-F])\)?[、\.．\s]\s*(.*)', re.IGNORECASE)
RE_ANSWER = re.compile(r'^\s*(?:正确)?答案[:：]\s*(.*)', re.IGNORECASE)
RE_ANALYSIS = re.compile(r'^\s*(?:答案)?解析[:：]\s*(.*)', re.IGNORECASE)
RE_CHAPTER_IN_ANA = re.compile(r'(?:<br>|[\s\n])*(所在章节[:：].*)')

# 记录日志的缓冲区
log_buffer = []

def log(msg):
    print(msg)
    log_buffer.append(msg)

def normalize_text(text):
    if not text: return ""
    text = re.sub(r'<[^>]+>', '', text) 
    return re.sub(r'[^\w\u4e00-\u9fa5]', '', text).strip()

def clean_answer(ans):
    if not ans: return ""
    ans = ans.strip().upper()
    if ans in ['√', 'T', 'Y', 'TRUE', '正确']: return '正确'
    if ans in ['×', 'F', 'N', 'FALSE', '错误']: return '错误'
    valid = [c for c in ans if c in "ABCDEF"]
    return "".join(valid)

def infer_type(filename):
    if '多选' in filename: return 'multi'
    if '单选' in filename: return 'single'
    if '判断' in filename: return 'tf'
    return 'unknown'

def parse_docx(file_path):
    fname = os.path.basename(file_path)
    try:
        doc = Document(file_path)
    except:
        log(f"❌ 无法读取文件: {fname}")
        return []

    questions = []
    curr = None
    collecting_txt = False
    
    # 统计该文件的行数
    total_paragraphs = len(doc.paragraphs)
    
    f_type = infer_type(fname)

    for i, para in enumerate(doc.paragraphs):
        line = para.text.strip()
        if not line: continue
        
        # 忽略页眉页脚噪音
        if "题库" in line and "竞赛" in line: continue
        if line.isdigit(): continue # 忽略纯数字页码

        # 1. 答案
        ans_match = RE_ANSWER.match(line)
        if ans_match:
            if curr:
                curr['answer'] = clean_answer(ans_match.group(1))
                collecting_txt = False
            continue
            
        # 2. 解析
        ana_match = RE_ANALYSIS.match(line)
        if ana_match:
            if curr:
                parts = RE_CHAPTER_IN_ANA.split(ana_match.group(1))
                if len(parts) > 1:
                    curr['analysis'] = parts[0].strip()
                    curr['chapter'] = parts[1].replace("所在章节：","").strip()
                else:
                    curr['analysis'] = parts[0].strip()
                collecting_txt = False
            continue

        # 3. 选项
        opt_match = RE_OPTION.match(line)
        if opt_match:
            if curr:
                if 'options' not in curr: curr['options'] = []
                curr['options'].append(line)
                collecting_txt = False
            continue

        # 4. 新题目判定
        is_new = False
        if curr is None: is_new = True
        elif 'answer' in curr and curr['answer']: 
            questions.append(curr)
            is_new = True
        
        if is_new:
            curr = {
                "question": line,
                "options": [],
                "answer": "",
                "analysis": "",
                "type": f_type,
                "source_file": fname,
                "line_no": i + 1 # 记录行号方便排查
            }
            collecting_txt = True
        elif collecting_txt and curr:
            curr['question'] += "<br>" + line

    if curr and 'answer' in curr:
        questions.append(curr)
    
    log(f"   📄 {fname}: 扫描 {total_paragraphs} 行 -> 提取 {len(questions)} 题")
    return questions

def quality_check(all_qs):
    """数据质量检查"""
    valid_qs = []
    invalid_qs = []
    
    for q in all_qs:
        reason = []
        # 检查1: 没有答案
        if not q['answer']:
            reason.append("缺失答案")
        
        # 检查2: 单/多选没有选项
        if q['type'] in ['single', 'multi'] and len(q['options']) < 2:
            reason.append("选项过少")
            
        # 检查3: 题干过短
        if len(normalize_text(q['question'])) < 3:
            reason.append("题干过短(疑似噪音)")
            
        if reason:
            q['error_reason'] = ", ".join(reason)
            invalid_qs.append(q)
        else:
            valid_qs.append(q)
            
    return valid_qs, invalid_qs

def process_and_deduplicate(all_qs):
    db = {"single": [], "multi": [], "tf": []}
    seen = {} 
    duplicates_info = [] # 记录重复详情
    
    for q in all_qs:
        # 类型兜底
        if q['type'] == 'unknown':
            if q['answer'] in ['正确', '错误']: q['type'] = 'tf'
            elif len(q['answer']) > 1: q['type'] = 'multi'
            else: q['type'] = 'single'
            
        if q['type'] == 'tf' and not q['options']:
            q['options'] = ["A. 正确", "B. 错误"]

        # 指纹
        fingerprint = normalize_text(q['question']) + q['answer']
        q_hash = hashlib.md5(fingerprint.encode('utf-8')).hexdigest()
        
        if q_hash in seen:
            old_q = seen[q_hash]
            # 记录重复信息
            duplicates_info.append(f"【重复】{q['source_file']} (行{q['line_no']}) 与 {old_q['source_file']} (行{old_q['line_no']}) 内容相同。保留解析较长者。")
            
            if len(q['analysis']) > len(old_q['analysis']):
                seen[q_hash] = q
        else:
            seen[q_hash] = q
            
    # 归类
    for q_hash, q in seen.items():
        q['id'] = q_hash 
        # 清理临时字段
        del q['source_file']
        del q['line_no']
        
        if q['type'] in db:
            db[q['type']].append(q)
            
    return db, duplicates_info

if __name__ == "__main__":
    log("=== 开始执行题库提取与验证程序 ===")
    
    if not os.path.exists(INPUT_DIR):
        log(f"❌ 找不到 {INPUT_DIR} 文件夹")
    else:
        # 1. 读取
        raw_data = []
        files = [f for f in os.listdir(INPUT_DIR) if f.endswith('.docx') and not f.startswith('~$')]
        
        if not files:
            log(f"⚠️  {INPUT_DIR} 中没有 .docx 文件")
        else:
            for f in files:
                raw_data.extend(parse_docx(os.path.join(INPUT_DIR, f)))
            
            log(f"\n📥 共提取原始题目: {len(raw_data)} 道")
            
            # 2. 质量检查
            valid_data, invalid_data = quality_check(raw_data)
            
            if invalid_data:
                log(f"⚠️  发现 {len(invalid_data)} 道异常题目 (将被丢弃，详情见报告底端)")
            else:
                log("✅ 数据质量检查通过，未发现格式错误的题目。")

            # 3. 去重
            final_db, dup_info = process_and_deduplicate(valid_data)
            
            # 4. 统计
            count_single = len(final_db['single'])
            count_multi = len(final_db['multi'])
            count_tf = len(final_db['tf'])
            total_final = count_single + count_multi + count_tf
            
            log(f"\n🧹 去重后有效题目: {total_final} 道")
            log(f"   ├─ 单选题: {count_single}")
            log(f"   ├─ 多选题: {count_multi}")
            log(f"   └─ 判断题: {count_tf}")
            
            # 5. 输出 JS
            js_content = f"const QUESTION_DB = {json.dumps(final_db, ensure_ascii=False, indent=2)};"
            with open(OUTPUT_JS, 'w', encoding='utf-8') as f:
                f.write(js_content)
            log(f"\n💾 题库文件已生成: {OUTPUT_JS}")
            
            # 6. 生成详细报告
            with open(REPORT_FILE, 'w', encoding='utf-8') as f:
                f.write("\n".join(log_buffer))
                f.write("\n\n" + "="*30 + " 重复题目明细 " + "="*30 + "\n")
                if dup_info:
                    f.write("\n".join(dup_info))
                else:
                    f.write("无重复题目。")
                
                f.write("\n\n" + "="*30 + " 异常/丢弃题目明细 " + "="*30 + "\n")
                if invalid_data:
                    for q in invalid_data:
                        f.write(f"❌ [{q['error_reason']}] {q['source_file']} (行{q['line_no']}): {q['question'][:30]}...\n")
                else:
                    f.write("无异常题目。")
                    
            print(f"📋 详细验证报告已生成: {os.path.abspath(REPORT_FILE)}")
            print("👉 请打开报告查看是否有题目遗漏！")